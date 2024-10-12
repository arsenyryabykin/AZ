from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
import sqlite3

# Создание документа по исходной конфигурации стелажей
def make_shelves_doc():
    conn = sqlite3.connect("imitators.db")
    cur = conn.cursor()
    cur.execute('''SELECT id, shelf_number, coordinate, tvs_number, suz_number
                    FROM shells''')
    data = cur.fetchall()

    document = Document()
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_cells = table.rows[0].cells
    headers_cells[0].text = '№'
    headers_cells[1].text = 'Номер стеллажа'
    headers_cells[2].text = 'Координата ячейки'
    headers_cells[3].text = 'Номер ИТВС'
    headers_cells[4].text = 'Номер ИПС СУЗ'
    for element in data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(element[0])
        row_cells[1].text = str(element[1])
        row_cells[2].text = element[2]
        row_cells[3].text = element[3]
        row_cells[4].text = element[4] if(element[4] != None) else ''

    document.save('Стеллажи.docx')
    print("Создание документа \"Стеллажи.docx\"")

# Создание документа по перегрузочным операциям Стеллаж -> Чехол
def make_case_operations_doc():
    conn = sqlite3.connect("imitators.db")
    cur = conn.cursor()
    cur.execute('''SELECT operation_id, shelf_number_from, shelf_coordinate_from,
                            case_number_to, case_coordinate_to, 
                            tvs_number, suz_number
                     FROM case_operations''')
    data = cur.fetchall()

    document = Document()
    table = document.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    headers_cells = table.rows[0].cells
    headers_cells[0].text = '№'
    headers_cells[1].text = 'Номер стеллажа'
    headers_cells[2].text = 'Координата ячейки стеллажа'
    headers_cells[3].text = 'Номер чехла'
    headers_cells[4].text = 'Координата ячейки чехла'
    headers_cells[5].text = 'Номер ИТВС'
    headers_cells[6].text = 'Номер ИПС СУЗ'
    for element in data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(element[0])
        row_cells[1].text = str(element[1])
        row_cells[2].text = element[2]
        row_cells[3].text = str(element[3])
        row_cells[4].text = element[4]
        row_cells[5].text = element[5]
        row_cells[6].text = element[6] if (element[6] != None) else ''

    document.save('Перекладывание в чехлы.docx')
    print("Создание документа \"Перекладывание в чехлы.docx\"")

# Создание документа по перегрузочным операциям Чехол -> АЗ
def make_az_operations_doc():
    conn = sqlite3.connect("imitators.db")
    cur = conn.cursor()
    cur.execute('''SELECT operation_id, case_number_from, case_coordinate_from,
                          az_coordinate_to, tvs_number, suz_number
                          FROM az_operations''')
    data = cur.fetchall()

    document = Document()
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    headers_cells = table.rows[0].cells
    headers_cells[0].text = '№ операции'
    headers_cells[1].text = 'Номер стеллажа'
    headers_cells[2].text = 'Координата ячейки стеллажа'
    headers_cells[3].text = 'Координата ячейки активной зоны'
    headers_cells[4].text = 'Номер ИТВС'
    headers_cells[5].text = 'Номер ИПС СУЗ'
    for element in data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(element[0])
        row_cells[1].text = str(element[1])
        row_cells[2].text = element[2]
        row_cells[3].text = element[3]
        row_cells[4].text = element[4]
        row_cells[5].text = element[5] if (element[5] != None) else ''

    document.save('Перекладывание в АЗ.docx')
    print("Создание документа \"Перекладывание в АЗ.docx\"")
